# backend/tools/docx_tools.py
from docx import Document
import io


def extract_docx(path: str) -> str:
    doc = Document(path)
    return "\n".join(para.text for para in doc.paragraphs)


def write_docx(sections: dict, output_path: str, ref_key: str = "references"):
    """
    Write formatted sections to a .docx file.
    ref_key: 'references' for APA, 'works_cited' for MLA, 
             'bibliography' for Chicago — passed in from output_node.
    """
    doc = Document()

    # Title
    if sections.get("title"):
        doc.add_heading(sections["title"], level=0)

    # Authors
    if sections.get("authors"):
        authors = sections["authors"]
        if isinstance(authors, list):
            authors = ", ".join(authors)
        doc.add_paragraph(authors)

    # First page header (MLA)
    fph = sections.get("first_page_header")
    if fph and any(fph.values()):
        for field in ["author", "instructor", "course", "date"]:
            if fph.get(field):
                doc.add_paragraph(fph[field])

    # Abstract (APA)
    if sections.get("abstract"):
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(sections["abstract"])
        if sections.get("keywords"):
            kw = sections["keywords"]
            if isinstance(kw, list):
                kw = ", ".join(kw)
            doc.add_paragraph(f"Keywords: {kw}")

    # Body sections
    for heading, content in sections.get("body", {}).items():
        doc.add_heading(heading, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph(str(content))

    # Footnotes (Chicago NB)
    if sections.get("footnotes"):
        doc.add_heading("Notes", level=1)
        for fn in sections["footnotes"]:
            doc.add_paragraph(str(fn))

    # Reference list — generic: works_cited (MLA), references (APA), bibliography (Chicago)
    ref_list = sections.get(ref_key) or sections.get("references") or sections.get("works_cited")
    if ref_list:
        # Label: "Works Cited" for MLA, "References" for APA/Chicago AD, "Bibliography" for Chicago NB
        label_map = {
            "works_cited":  "Works Cited",
            "references":   "References",
            "bibliography": "Bibliography",
        }
        label = label_map.get(ref_key, "References")
        doc.add_heading(label, level=1)
        for ref in ref_list:
            doc.add_paragraph(str(ref))

    doc.save(output_path)


def write_docx_to_buffer(sections: dict, ref_key: str = "references") -> io.BytesIO:
    """
    Same as write_docx but writes to an in-memory BytesIO buffer
    instead of a file, so the API can stream it to the browser without
    touching the filesystem.
    """
    doc = Document()

    if sections.get("title"):
        doc.add_heading(sections["title"], level=0)

    if sections.get("authors"):
        authors = sections["authors"]
        if isinstance(authors, list):
            authors = ", ".join(authors)
        doc.add_paragraph(authors)

    fph = sections.get("first_page_header")
    if fph and any(fph.values()):
        for field in ["author", "instructor", "course", "date"]:
            if fph.get(field):
                doc.add_paragraph(fph[field])

    if sections.get("abstract"):
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(sections["abstract"])
        if sections.get("keywords"):
            kw = sections["keywords"]
            if isinstance(kw, list):
                kw = ", ".join(kw)
            doc.add_paragraph(f"Keywords: {kw}")

    for heading, content in sections.get("body", {}).items():
        doc.add_heading(heading, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(str(item))
        else:
            doc.add_paragraph(str(content))

    if sections.get("footnotes"):
        doc.add_heading("Notes", level=1)
        for fn in sections["footnotes"]:
            doc.add_paragraph(str(fn))

    ref_list = sections.get(ref_key) or sections.get("references") or sections.get("works_cited")
    if ref_list:
        label_map = {
            "works_cited":  "Works Cited",
            "references":   "References",
            "bibliography": "Bibliography",
        }
        label = label_map.get(ref_key, "References")
        doc.add_heading(label, level=1)
        for ref in ref_list:
            doc.add_paragraph(str(ref))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
